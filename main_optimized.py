import streamlit as st
import os
from dotenv import load_dotenv
import pandas as pd
from anthropic import Anthropic
from datetime import datetime
import hashlib
import io
import base64
import tempfile

# Import untuk handle berbagai file types
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# OCR Libraries untuk scanned PDFs
try:
    import pytesseract
    from PIL import Image
    import pdf2image
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

try:
    import easyocr
    HAS_EASYOCR = True
except ImportError:
    HAS_EASYOCR = False

# Load environment variables
load_dotenv()

st.title("🤖 AI Assistant with OCR Support")
st.caption("Handles scanned PDFs, images, and all document types")

# Initialize Anthropic client
api_key = os.getenv("ANTHROPIC_API_KEY")
if api_key:
    client = Anthropic(api_key=api_key)
else:
    client = None
    st.error("Please set ANTHROPIC_API_KEY in .env file")

# Initialize OCR reader if available
@st.cache_resource
def init_ocr_reader():
    """Initialize OCR reader (cached for performance)"""
    if HAS_EASYOCR:
        try:
            import easyocr
            reader = easyocr.Reader(['en', 'id'])  # English + Indonesian
            return reader
        except:
            pass
    return None

ocr_reader = init_ocr_reader()

# Constants
MAX_CONTEXT_TOKENS = 45000
TOKENS_PER_CHAR = 0.25

# Roles
ROLES = {
    "General Assistant": {
        "system_prompt": "You are a helpful AI assistant. Be friendly, informative, and professional.",
        "icon": "🤖",
    },
    "Human Resources": {
        "system_prompt": """You are an HR specialist focused on data analysis and promotion evaluations. You should:
        - Analyze employee performance data objectively and thoroughly
        - Consider all metrics: experience, skills, achievements, performance scores
        - Provide specific, data-driven recommendations with citations
        - Compare employees fairly based on quantitative metrics
        - Identify top performers and areas for improvement
        - Use exact numbers and percentages from the data
        - Make promotion recommendations based on clear criteria""",
        "icon": "📊",
    },
    "Data Analyst": {
        "system_prompt": """You are a data analyst expert. You should:
        - Perform thorough statistical analysis on all data
        - Identify trends, patterns, and outliers
        - Provide specific numbers and percentages
        - Create clear summaries of complex datasets
        - Make data-driven recommendations
        - Always cite exact figures from the source""",
        "icon": "📈",
    },
}

# ========== OCR FUNCTIONS FOR SCANNED PDFS ==========

def is_scanned_pdf(pdf_file):
    """Check if PDF is scanned (image-based) or has extractable text"""
    has_text = False
    
    if HAS_PYMUPDF:
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(pdf_file.getvalue())
                tmp_path = tmp.name
            
            doc = fitz.open(tmp_path)
            # Check first few pages for text
            for page_num in range(min(3, len(doc))):
                page = doc.load_page(page_num)
                text = page.get_text().strip()
                if len(text) > 50:  # If substantial text found
                    has_text = True
                    break
            doc.close()
            os.unlink(tmp_path)
        except:
            pass
    
    elif HAS_PYPDF2:
        try:
            pdf_file.seek(0)
            reader = PyPDF2.PdfReader(pdf_file)
            # Check first few pages
            for page_num in range(min(3, len(reader.pages))):
                text = reader.pages[page_num].extract_text().strip()
                if len(text) > 50:
                    has_text = True
                    break
        except:
            pass
    
    return not has_text  # Return True if scanned (no text found)

def extract_text_with_tesseract(pdf_file):
    """Extract text from scanned PDF using Tesseract OCR"""
    if not HAS_TESSERACT:
        return None, "Tesseract not installed"
    
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(pdf_file.getvalue())
            tmp_path = tmp.name
        
        # Convert PDF to images
        images = pdf2image.convert_from_path(tmp_path, dpi=200)
        
        all_text = ""
        progress = st.progress(0)
        
        for i, image in enumerate(images):
            # Update progress
            progress.progress((i + 1) / len(images), f"OCR Page {i + 1}/{len(images)}")
            
            # OCR with language support
            try:
                # Try with Indonesian + English
                text = pytesseract.image_to_string(image, lang='ind+eng')
            except:
                # Fallback to English only
                text = pytesseract.image_to_string(image, lang='eng')
            
            all_text += f"\n--- Page {i + 1} ---\n{text}\n"
        
        progress.empty()
        os.unlink(tmp_path)
        
        return all_text, "success"
        
    except Exception as e:
        return None, f"Tesseract OCR failed: {str(e)}"

def extract_text_with_easyocr(pdf_file):
    """Extract text from scanned PDF using EasyOCR"""
    if not HAS_EASYOCR or not ocr_reader:
        return None, "EasyOCR not available"
    
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(pdf_file.getvalue())
            tmp_path = tmp.name
        
        # Convert PDF to images
        images = pdf2image.convert_from_path(tmp_path, dpi=200)
        
        all_text = ""
        progress = st.progress(0)
        
        for i, image in enumerate(images):
            # Update progress
            progress.progress((i + 1) / len(images), f"OCR Page {i + 1}/{len(images)}")
            
            # Convert PIL to numpy array for EasyOCR
            import numpy as np
            img_array = np.array(image)
            
            # Run OCR
            results = ocr_reader.readtext(img_array)
            
            # Extract text from results
            page_text = ""
            for (bbox, text, prob) in results:
                if prob > 0.5:  # Only include confident detections
                    page_text += text + " "
            
            all_text += f"\n--- Page {i + 1} ---\n{page_text}\n"
        
        progress.empty()
        os.unlink(tmp_path)
        
        return all_text, "success"
        
    except Exception as e:
        return None, f"EasyOCR failed: {str(e)}"

def extract_text_from_pdf(pdf_file, force_ocr=False):
    """Extract text from PDF with automatic detection of scanned PDFs"""
    
    # First, check if it's a scanned PDF
    if force_ocr or is_scanned_pdf(pdf_file):
        st.info("📸 Detected scanned PDF. Using OCR...")
        
        # Try different OCR methods
        if HAS_TESSERACT:
            text, status = extract_text_with_tesseract(pdf_file)
            if text:
                st.success("✅ OCR completed with Tesseract")
                return text
        
        if HAS_EASYOCR:
            text, status = extract_text_with_easyocr(pdf_file)
            if text:
                st.success("✅ OCR completed with EasyOCR")
                return text
        
        # No OCR available
        st.error("❌ This appears to be a scanned PDF but no OCR libraries are installed")
        with st.expander("📚 Install OCR Support"):
            st.markdown("""
            **Option 1: Tesseract (Recommended)**
            ```bash
            # Install Tesseract
            # Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki
            # Mac: brew install tesseract
            # Linux: sudo apt-get install tesseract-ocr
            
            # Install Python packages
            pip install pytesseract pillow pdf2image
            
            # For Indonesian support
            # Windows: Install language pack during Tesseract installation
            # Linux: sudo apt-get install tesseract-ocr-ind
            ```
            
            **Option 2: EasyOCR (Easier setup, slower)**
            ```bash
            pip install easyocr pdf2image
            ```
            """)
        return None
    
    # Regular text extraction for non-scanned PDFs
    text = ""
    
    # Try PyMuPDF first
    if HAS_PYMUPDF:
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(pdf_file.getvalue())
                tmp_path = tmp.name
            
            doc = fitz.open(tmp_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page.get_text()
            doc.close()
            os.unlink(tmp_path)
            
            if text.strip():
                return text
        except Exception as e:
            st.warning(f"PyMuPDF extraction failed: {e}")
    
    # Fallback to PyPDF2
    if HAS_PYPDF2:
        try:
            pdf_file.seek(0)
            reader = PyPDF2.PdfReader(pdf_file)
            for page_num, page in enumerate(reader.pages):
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page.extract_text()
            
            if text.strip():
                return text
        except Exception as e:
            st.error(f"PDF extraction failed: {e}")
    
    return None

# ========== OTHER EXTRACTION FUNCTIONS ==========

def extract_text_from_excel(excel_file):
    """Extract and format Excel data with comprehensive analysis"""
    try:
        excel_data = pd.ExcelFile(excel_file)
        all_text = f"# EXCEL DOCUMENT: {excel_file.name}\n"
        all_text += f"Total Sheets: {len(excel_data.sheet_names)}\n"
        all_text += "="*60 + "\n\n"
        
        for sheet_name in excel_data.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            
            all_text += f"\n## SHEET: {sheet_name}\n"
            all_text += f"Dimensions: {df.shape[0]} rows × {df.shape[1]} columns\n"
            all_text += f"Columns: {', '.join(str(col) for col in df.columns)}\n"
            all_text += "-"*40 + "\n\n"
            
            # Data representation
            if len(df) <= 100:
                all_text += "### COMPLETE DATA:\n```\n"
                all_text += df.to_string(max_rows=None, max_cols=None)
                all_text += "\n```\n\n"
            else:
                all_text += "### DATA SAMPLE (First 50 rows):\n```\n"
                all_text += df.head(50).to_string(max_cols=None)
                all_text += f"\n```\n... ({len(df) - 50} more rows) ...\n\n"
            
            # Statistics
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                all_text += "### STATISTICAL ANALYSIS:\n```\n"
                all_text += df[numeric_cols].describe().to_string()
                all_text += "\n```\n\n"
            
            # Categorical analysis
            categorical_cols = df.select_dtypes(include=['object']).columns
            for col in categorical_cols[:5]:
                if df[col].nunique() <= 20:
                    all_text += f"\n### {col} DISTRIBUTION:\n```\n"
                    all_text += df[col].value_counts().to_string()
                    all_text += "\n```\n\n"
        
        return all_text
    except Exception as e:
        st.error(f"Excel extraction error: {str(e)}")
        return None

def extract_text_from_csv(csv_file):
    """Extract and format CSV data"""
    try:
        for encoding in ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']:
            try:
                csv_file.seek(0)
                df = pd.read_csv(csv_file, encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            return None
        
        all_text = f"# CSV DOCUMENT: {csv_file.name}\n"
        all_text += f"Dimensions: {df.shape[0]} rows × {df.shape[1]} columns\n"
        all_text += "="*60 + "\n\n"
        
        if len(df) <= 100:
            all_text += "## COMPLETE DATA:\n```\n"
            all_text += df.to_string(max_rows=None, max_cols=None)
            all_text += "\n```\n\n"
        else:
            all_text += "## DATA SAMPLE:\n```\n"
            all_text += df.head(50).to_string(max_cols=None)
            all_text += f"\n```\n... ({len(df) - 50} more rows) ...\n\n"
        
        # Statistics
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            all_text += "## STATISTICS:\n```\n"
            all_text += df[numeric_cols].describe().to_string()
            all_text += "\n```\n"
        
        return all_text
    except Exception as e:
        st.error(f"CSV extraction error: {str(e)}")
        return None

def extract_text_from_word(word_file):
    """Extract text from Word document"""
    if not HAS_DOCX:
        st.error("python-docx not installed")
        return None
    
    try:
        doc = docx.Document(word_file)
        all_text = f"# WORD DOCUMENT: {word_file.name}\n{'='*60}\n\n"
        
        for para in doc.paragraphs:
            if para.text.strip():
                all_text += para.text + "\n\n"
        
        if doc.tables:
            all_text += "\n## TABLES:\n\n"
            for i, table in enumerate(doc.tables):
                all_text += f"### Table {i+1}:\n"
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    all_text += " | ".join(row_data) + "\n"
                all_text += "\n"
        
        return all_text
    except Exception as e:
        st.error(f"Word extraction error: {str(e)}")
        return None

def extract_text_from_txt(txt_file):
    """Extract text from plain text file"""
    try:
        for encoding in ['utf-8', 'latin-1', 'iso-8859-1']:
            try:
                txt_file.seek(0)
                return f"# TEXT DOCUMENT: {txt_file.name}\n{'='*60}\n\n" + txt_file.read().decode(encoding)
            except UnicodeDecodeError:
                continue
        return None
    except:
        return None

def extract_text_from_image(image_file):
    """Extract text from image files using OCR"""
    if not HAS_TESSERACT and not HAS_EASYOCR:
        st.error("No OCR libraries installed for image text extraction")
        return None
    
    try:
        # Load image
        from PIL import Image
        image = Image.open(image_file)
        
        # Try Tesseract first
        if HAS_TESSERACT:
            try:
                text = pytesseract.image_to_string(image, lang='ind+eng')
                if text.strip():
                    return f"# IMAGE OCR: {image_file.name}\n{'='*60}\n\n{text}"
            except:
                pass
        
        # Try EasyOCR
        if HAS_EASYOCR and ocr_reader:
            try:
                import numpy as np
                results = ocr_reader.readtext(np.array(image))
                text = " ".join([text for _, text, prob in results if prob > 0.5])
                if text.strip():
                    return f"# IMAGE OCR: {image_file.name}\n{'='*60}\n\n{text}"
            except:
                pass
        
        st.error("OCR extraction failed for image")
        return None
        
    except Exception as e:
        st.error(f"Image processing error: {str(e)}")
        return None

# ========== DOCUMENT MANAGER ==========

class DocumentManager:
    def __init__(self):
        self.documents = {}
        self.context = ""
        self.total_tokens = 0
    
    def add_document(self, name, content, file_type):
        """Add document to manager"""
        if not content:
            return None, 0
        
        doc_id = hashlib.md5(f"{name}{len(content)}".encode()).hexdigest()[:8]
        tokens = int(len(content) * TOKENS_PER_CHAR)
        
        self.documents[doc_id] = {
            'name': name,
            'content': content,
            'type': file_type,
            'tokens': tokens,
            'added': datetime.now()
        }
        
        self.update_context()
        return doc_id, tokens
    
    def update_context(self):
        """Update combined context"""
        parts = []
        tokens = 0
        
        # Sort by size to fit more
        sorted_docs = sorted(self.documents.items(), key=lambda x: x[1]['tokens'])
        
        for doc_id, doc in sorted_docs:
            if tokens + doc['tokens'] <= MAX_CONTEXT_TOKENS:
                parts.append(f"\n{'='*60}\nDOCUMENT: {doc['name']}\n{'='*60}\n{doc['content']}")
                tokens += doc['tokens']
            elif tokens < MAX_CONTEXT_TOKENS * 0.9:
                # Add partial if space
                remaining = int((MAX_CONTEXT_TOKENS - tokens) / TOKENS_PER_CHAR)
                if remaining > 1000:
                    partial = doc['content'][:remaining]
                    parts.append(f"\n{'='*60}\nDOCUMENT: {doc['name']} [TRUNCATED]\n{'='*60}\n{partial}\n...[truncated]")
                    break
        
        self.context = '\n'.join(parts)
        self.total_tokens = tokens
    
    def remove_document(self, doc_id):
        """Remove document"""
        if doc_id in self.documents:
            del self.documents[doc_id]
            self.update_context()
    
    def clear_all(self):
        """Clear all documents"""
        self.documents = {}
        self.context = ""
        self.total_tokens = 0

# ========== STREAMLIT APP ==========

# Initialize
if 'doc_manager' not in st.session_state:
    st.session_state.doc_manager = DocumentManager()

if 'messages' not in st.session_state:
    st.session_state.messages = []

# Sidebar
with st.sidebar:
    st.header("⚙️ Configuration")
    
    # Status
    st.subheader("🔌 System Status")
    col1, col2 = st.columns(2)
    
    with col1:
        if client:
            st.success("✅ API Key")
        else:
            st.error("❌ API Key")
    
    with col2:
        if HAS_TESSERACT or HAS_EASYOCR:
            st.success("✅ OCR")
        else:
            st.warning("⚠️ No OCR")
    
    # Libraries status
    with st.expander("📚 Libraries Status"):
        libs = {
            "PyPDF2": ("PDF text extraction", HAS_PYPDF2),
            "PyMuPDF": ("Better PDF support", HAS_PYMUPDF),
            "python-docx": ("Word documents", HAS_DOCX),
            "Tesseract": ("OCR for scans", HAS_TESSERACT),
            "EasyOCR": ("Alternative OCR", HAS_EASYOCR)
        }
        
        for lib, (desc, status) in libs.items():
            if status:
                st.success(f"✅ {lib}: {desc}")
            else:
                st.warning(f"⚠️ {lib}: {desc}")
        
        if not HAS_TESSERACT and not HAS_EASYOCR:
            st.error("⚠️ No OCR support - cannot process scanned PDFs!")
            if st.button("📖 See OCR Installation Guide"):
                st.markdown("""
                **Tesseract (Recommended):**
                ```bash
                # Windows
                Download from: https://github.com/UB-Mannheim/tesseract/wiki
                
                # Mac
                brew install tesseract
                brew install tesseract-lang  # All languages
                
                # Linux
                sudo apt-get install tesseract-ocr
                sudo apt-get install tesseract-ocr-ind  # Indonesian
                
                # Python packages
                pip install pytesseract pillow pdf2image
                ```
                
                **EasyOCR (Alternative):**
                ```bash
                pip install easyocr pdf2image
                ```
                """)
    
    # Model selection
    st.subheader("🧠 Model")
    model = st.selectbox(
        "Choose model:",
        ["claude-3-5-sonnet-20241022", "claude-3-opus-20240229",
         "claude-3-sonnet-20240229", "claude-3-haiku-20240307"]
    )
    
    # Role
    st.subheader("🎭 Role")
    role = st.selectbox(
        "Choose role:",
        list(ROLES.keys()),
        index=1
    )
    
    # Temperature
    temperature = st.slider("Temperature:", 0.0, 1.0, 0.3, 0.1)
    
    # File upload
    st.subheader("📚 Document Upload")
    
    uploaded = st.file_uploader(
        "Upload files:",
        type=["pdf", "txt", "xlsx", "xls", "csv", "docx", "doc", "png", "jpg", "jpeg"],
        accept_multiple_files=True,
        help="Supports scanned PDFs, images, and all document types"
    )
    
    # OCR Options for PDFs
    if any(f.name.endswith('.pdf') for f in uploaded) if uploaded else False:
        force_ocr = st.checkbox("🔍 Force OCR for all PDFs", help="Use OCR even if text is extractable")
    else:
        force_ocr = False
    
    if uploaded and client:
        for file in uploaded:
            # Check if already loaded
            exists = any(doc['name'] == file.name for doc in st.session_state.doc_manager.documents.values())
            
            if not exists:
                ext = file.name.split('.')[-1].lower()
                
                with st.spinner(f"Processing {file.name}..."):
                    text = None
                    
                    # Process based on file type
                    if ext == 'pdf':
                        text = extract_text_from_pdf(file, force_ocr=force_ocr)
                    elif ext in ['xlsx', 'xls']:
                        text = extract_text_from_excel(file)
                    elif ext == 'csv':
                        text = extract_text_from_csv(file)
                    elif ext in ['docx', 'doc']:
                        text = extract_text_from_word(file)
                    elif ext == 'txt':
                        text = extract_text_from_txt(file)
                    elif ext in ['png', 'jpg', 'jpeg']:
                        text = extract_text_from_image(file)
                    
                    if text:
                        doc_id, tokens = st.session_state.doc_manager.add_document(
                            file.name, text, ext
                        )
                        if doc_id:
                            icon = {
                                'pdf': '📕', 'xlsx': '📊', 'xls': '📊',
                                'csv': '📈', 'docx': '📄', 'doc': '📄',
                                'txt': '📝', 'png': '🖼️', 'jpg': '🖼️', 'jpeg': '🖼️'
                            }.get(ext, '📎')
                            st.success(f"{icon} {file.name} → {tokens:,} tokens")
                    else:
                        st.error(f"❌ Failed to process {file.name}")
    
    # Display documents
    if st.session_state.doc_manager.documents:
        st.divider()
        st.write("**Loaded Documents:**")
        
        for doc_id, doc in st.session_state.doc_manager.documents.items():
            col1, col2 = st.columns([4, 1])
            with col1:
                icon = {
                    'pdf': '📕', 'xlsx': '📊', 'xls': '📊',
                    'csv': '📈', 'docx': '📄', 'doc': '📄',
                    'txt': '📝', 'png': '🖼️', 'jpg': '🖼️', 'jpeg': '🖼️'
                }.get(doc['type'], '📎')
                st.caption(f"{icon} {doc['name']} ({doc['tokens']:,} tokens)")
            with col2:
                if st.button("🗑️", key=f"del_{doc_id}"):
                    st.session_state.doc_manager.remove_document(doc_id)
                    st.rerun()
        
        # Token usage
        st.divider()
        tokens = st.session_state.doc_manager.total_tokens
        percent = (tokens / MAX_CONTEXT_TOKENS) * 100
        
        if percent < 80:
            st.progress(percent / 100, f"{tokens:,} / {MAX_CONTEXT_TOKENS:,} tokens")
        else:
            st.warning(f"⚠️ {tokens:,} / {MAX_CONTEXT_TOKENS:,} tokens")
        
        # Clear all
        if st.button("🗑️ Clear All Documents"):
            st.session_state.doc_manager.clear_all()
            st.session_state.messages = []
            st.rerun()

# Main interface
col1, col2, col3 = st.columns(3)
with col1:
    st.markdown(f"**Role:** {ROLES[role]['icon']} {role}")
with col2:
    st.markdown(f"**Model:** {model.split('-')[1].title()}")
with col3:
    if st.session_state.doc_manager.documents:
        st.markdown(f"**Docs:** {len(st.session_state.doc_manager.documents)}")

# OCR indicator
if any(doc['type'] == 'pdf' for doc in st.session_state.doc_manager.documents.values()):
    if HAS_TESSERACT or HAS_EASYOCR:
        st.success("✅ OCR ready for scanned documents")
    else:
        st.warning("⚠️ OCR not available - scanned PDFs may not be readable")

# Chat history
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# Chat input
if prompt := st.chat_input("Ask about your documents..."):
    if not client:
        st.error("Please configure API key")
    else:
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        with st.chat_message("user"):
            st.markdown(prompt)
        
        with st.chat_message("assistant"):
            try:
                # Build system prompt with context
                system = ROLES[role]["system_prompt"]
                
                if st.session_state.doc_manager.context:
                    system += f"""

You have access to the following documents:

{st.session_state.doc_manager.context}

IMPORTANT: Use specific data and numbers from these documents when answering."""
                
                # Format messages
                messages = [
                    {"role": "user" if m["role"] == "user" else "assistant", "content": m["content"]}
                    for m in st.session_state.messages
                ]
                
                # Stream response
                placeholder = st.empty()
                full_response = ""
                
                with client.messages.stream(
                    model=model,
                    max_tokens=4096,
                    temperature=temperature,
                    system=system,
                    messages=messages
                ) as stream:
                    for text in stream.text_stream:
                        full_response += text
                        placeholder.markdown(full_response + "▌")
                
                placeholder.markdown(full_response)
                st.session_state.messages.append({"role": "assistant", "content": full_response})
                
            except Exception as e:
                st.error(f"Error: {str(e)}")

# Clear chat
if st.button("🔄 Clear Chat"):
    st.session_state.messages = []
    st.rerun()

# Help
with st.expander("ℹ️ OCR Support & Features"):
    st.markdown("""
    ### Complete Document Processing with OCR
    
    **Supported File Types:**
    - 📕 **PDF** - Both text and scanned (with OCR)
    - 🖼️ **Images** - PNG, JPG, JPEG (with OCR)
    - 📊 **Excel** - Full statistical analysis
    - 📈 **CSV** - Complete data processing
    - 📄 **Word** - Text and tables
    - 📝 **Text** - Plain text files
    
    **OCR Features:**
    - ✅ **Automatic detection** of scanned PDFs
    - ✅ **Multi-language** support (English + Indonesian)
    - ✅ **Progress tracking** during OCR
    - ✅ **Multiple OCR engines** (Tesseract, EasyOCR)
    - ✅ **Force OCR option** for mixed PDFs
    
    **Installation for OCR Support:**
    
    **Option 1: Tesseract (Best quality)**
    ```bash
    # Windows
    Download installer from:
    https://github.com/UB-Mannheim/tesseract/wiki
    
    # Mac
    brew install tesseract
    brew install tesseract-lang
    
    # Ubuntu/Debian
    sudo apt-get install tesseract-ocr
    sudo apt-get install tesseract-ocr-ind
    
    # Python packages
    pip install pytesseract pillow pdf2image
    ```
    
    **Option 2: EasyOCR (Easier setup)**
    ```bash
    pip install easyocr pdf2image
    ```
    
    **Option 3: For poppler (pdf2image dependency)**
    ```bash
    # Windows
    Download from: https://github.com/oschwartz10612/poppler-windows/releases
    Add to PATH
    
    # Mac
    brew install poppler
    
    # Linux
    sudo apt-get install poppler-utils
    ```
    
    **Tips:**
    - For best OCR results, use high-quality scans (200+ DPI)
    - Check "Force OCR" for PDFs with mixed text/images
    - Indonesian documents work best with Tesseract + ind language pack
    """)